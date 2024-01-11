from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pystrich.datamatrix import DataMatrixEncoder
from io import BytesIO
from PIL import Image


def set_font(paragraph, font_name, font_size):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = True


def create_barcode_image(data):
    encoder = DataMatrixEncoder(data)
    encoder.save('barcode.png')

    with Image.open('barcode.png') as img:
        bbox = img.getbbox()
        cropped_img = img.crop(bbox)
        img_byte_arr = BytesIO()
        cropped_img.save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)
        return img_byte_arr


def trim_whitespace(image_stream, border_size=10):
    image = Image.open(image_stream)
    bbox = image.getbbox()

    # Reduce the bounding box by the border size, effectively cropping the border
    cropped_bbox = (bbox[0] + border_size, bbox[1] + border_size, bbox[2] - border_size, bbox[3] - border_size)
    cropped_image = image.crop(cropped_bbox)

    new_image_stream = BytesIO()
    cropped_image.save(new_image_stream, format="PNG")
    new_image_stream.seek(0)
    return new_image_stream


def create_datamatrix_doc_from_text(text_content):
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.3)
        section.bottom_margin = Cm(0.3)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.3)

    products = text_content.strip().split('\n\n')
    for product in products:
        lines = product.strip().split('\n')
        product_name = lines[0]
        codes = lines[1:]

        document.add_heading(product_name, level=1)

        for _ in range(12):
            codes_str = ""
            count = 0

            for code in codes:
                if count == 0 or count == 3 or count == 5 or count == 7:
                    spacer = " " * 9
                else:
                    spacer = " " * 10

                codes_str += code[1:14] + spacer
                count += 1

            paragraph_codes = document.add_paragraph(codes_str)
            paragraph_codes.paragraph_format.space_after = Pt(0)
            paragraph_codes.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            set_font(paragraph_codes, 'Arial Narrow', 5)

            paragraph_barcodes = document.add_paragraph()
            paragraph_codes.paragraph_format.space_after = Pt(0)
            paragraph_barcodes.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for code in codes:
                barcode_image_stream = create_barcode_image(code)
                trimmed_barcode_image_stream = trim_whitespace(barcode_image_stream)
                run = paragraph_barcodes.add_run()
                run.add_picture(trimmed_barcode_image_stream, width=Cm(1), height=Cm(1))
                run.add_text('     ')

        document.add_page_break()

    return document


with open('datamatrix-barcode-docx/text.txt', 'r') as file:
    text_content = file.read()

barcode_doc = create_datamatrix_doc_from_text(text_content)

output_path = 'all_products_barcodes.docx'
barcode_doc.save(output_path)
